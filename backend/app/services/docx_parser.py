"""
.docx -> raw text extraction.

Mirrors pdf_parser.py's contract (bytes in, text out, ValueError if empty)
so extraction.py can treat PDF and Word uploads identically once we're past
this stage.

Also handles .dotx (Word template) files. These are structurally identical
to .docx under the hood — the only difference python-docx cares about is
a single content-type string in [Content_Types].xml declaring the main
document part as a "template" rather than a "document". Rather than
requiring users to manually save-as .docx first (a real point of friction —
many real-world syllabi are distributed as .dotx template files, as
Drexel's are), we patch that string in memory before handing the bytes to
python-docx. This is NOT a general OOXML converter — it only fixes the one
content-type declaration that blocks python-docx from opening the file;
it doesn't handle .doc (legacy binary format) at all.
"""

import io
import zipfile

import docx


def _dotx_bytes_to_docx_bytes(file_bytes: bytes) -> bytes:
    src = zipfile.ZipFile(io.BytesIO(file_bytes))
    out_buffer = io.BytesIO()
    with zipfile.ZipFile(out_buffer, "w", zipfile.ZIP_DEFLATED) as out:
        for item in src.namelist():
            data = src.read(item)
            if item == "[Content_Types].xml":
                data = data.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            out.writestr(item, data)
    return out_buffer.getvalue()


def extract_text(file_bytes: bytes, is_template: bool = False) -> str:
    """
    Extracts text from a .docx (or .dotx, via is_template=True) file,
    including both regular paragraphs and table cells — syllabus grading
    breakdowns are often in actual Word tables, not paragraph text, so
    skipping tables would silently lose the most important data on the page.
    """
    if is_template:
        file_bytes = _dotx_bytes_to_docx_bytes(file_bytes)

    document = docx.Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                # Join cells with a colon-ish separator so "Midterm | 25%"
                # style table rows still match the "Name: XX%" regex pattern
                # used downstream.
                parts.append(": ".join(cell_texts))

    full_text = "\n".join(parts).strip()

    if not full_text:
        raise ValueError("No extractable text found in this Word document.")

    return full_text
